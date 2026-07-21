"""Export service.

Converts Markdown content to downloadable formats:
  - Markdown: raw UTF-8 bytes.
  - PDF: Markdown -> HTML -> WeasyPrint-rendered PDF.
  - DOCX: Markdown -> HTML -> python-docx with rich formatting.
"""

import io

import markdown as md_lib
from flask import current_app


_MD_EXTENSIONS = [
    'fenced_code',
    'tables',
    'sane_lists',
    'nl2br',
]


def export_as_markdown(content: str) -> bytes:
    return content.encode('utf-8')


def _markdown_to_html(content: str) -> str:
    return md_lib.markdown(content or '', extensions=_MD_EXTENSIONS)


_PDF_CSS = """
@page {
    size: A4 portrait;
    margin: 1.4cm;
}

@page wide {
    size: A4 landscape;
    margin: 1.2cm;
}

body {
    font-family: "Segoe UI", "Helvetica Neue", Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.55;
    color: #1a1a1a;
    overflow-wrap: anywhere;
}

h1, h2, h3, h4, h5, h6 {
    line-height: 1.3;
    margin-top: 0.8em;
    margin-bottom: 0.35em;
    break-after: avoid;
}

h1 { font-size: 19pt; }
h2 { font-size: 16pt; }
h3 { font-size: 14pt; }
h4, h5, h6 { font-size: 12pt; }

p { margin: 0.35em 0; }

pre {
    background: #f6f7f8;
    border: 1px solid #d9dde1;
    border-radius: 4px;
    padding: 10px;
    white-space: pre-wrap;
    overflow-wrap: anywhere;
    font-family: "Consolas", "Fira Code", monospace;
    font-size: 9.6pt;
}

code {
    font-family: "Consolas", "Fira Code", monospace;
    font-size: 9.4pt;
    background: #f1f3f5;
    border-radius: 3px;
    padding: 1px 4px;
}

pre code {
    background: transparent;
    padding: 0;
}

blockquote {
    border-left: 3px solid #c9cfd6;
    margin: 0.45em 0;
    padding: 0.25em 0.8em;
    background: #fafbfc;
    color: #434b53;
}

ul, ol {
    margin: 0.35em 0;
    padding-left: 1.4em;
}

li {
    margin-bottom: 0.15em;
}

hr {
    border: none;
    border-top: 1px solid #d9dde1;
    margin: 0.9em 0;
}

a {
    color: #0b57d0;
    text-decoration: none;
}

.table-wrap {
    margin: 0.6em 0 0.9em;
    break-inside: auto;
}

.table-wrap.wide {
    page: wide;
}

table {
    width: 100%;
    border-collapse: collapse;
    table-layout: fixed;
    font-size: 9pt;
}

table.wide {
    font-size: 8.4pt;
}

table.very-wide {
    font-size: 7.8pt;
}

thead {
    display: table-header-group;
}

tfoot {
    display: table-footer-group;
}

tr {
    break-inside: avoid;
}

th, td {
    border: 1px solid #c8ced4;
    padding: 4px 6px;
    vertical-align: top;
    text-align: left;
    overflow-wrap: anywhere;
    word-break: break-word;
    white-space: normal;
}

th {
    background: #edf1f5;
    font-weight: 600;
}
"""


def _prepare_pdf_html(html_body: str) -> str:
    try:
        from bs4 import BeautifulSoup
    except Exception:
        return html_body

    soup = BeautifulSoup(html_body, 'html.parser')

    for table in soup.find_all('table'):
        max_cols = 0
        for row in table.find_all('tr'):
            cols = 0
            for cell in row.find_all(['th', 'td']):
                try:
                    span = int(cell.get('colspan', 1))
                except (TypeError, ValueError):
                    span = 1
                cols += max(span, 1)
            max_cols = max(max_cols, cols)

        classes = table.get('class', [])
        if max_cols >= 12:
            classes.extend(['wide', 'very-wide'])
        elif max_cols >= 8:
            classes.append('wide')
        table['class'] = classes

        wrapper = soup.new_tag('div')
        wrapper_classes = ['table-wrap']
        if max_cols >= 12:
            wrapper_classes.append('wide')
        wrapper['class'] = wrapper_classes

        table.wrap(wrapper)

    return str(soup)


def export_as_pdf(content: str) -> bytes:
    html_body = _prepare_pdf_html(_markdown_to_html(content))

    html_doc = f"""<!DOCTYPE html>
<html lang=\"en\">
<head>
    <meta charset=\"utf-8\">
    <style>{_PDF_CSS}</style>
</head>
<body>
{html_body}
</body>
</html>
"""

    try:
        from weasyprint import HTML as WeasyprintHTML

        return WeasyprintHTML(string=html_doc).write_pdf()
    except ImportError:
        current_app.logger.warning('WeasyPrint is not installed; returning HTML bytes instead of PDF.')
        return html_doc.encode('utf-8')
    except Exception as exc:
        current_app.logger.error('WeasyPrint PDF generation failed: %s', exc)
        return html_doc.encode('utf-8')


def export_as_docx(content: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    try:
        from bs4 import BeautifulSoup, NavigableString, Tag
    except ImportError as exc:
        raise RuntimeError('beautifulsoup4 is required for DOCX export formatting.') from exc

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    def set_cell_shading(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tc_pr.append(shd)

    def add_inline(paragraph, node, bold=False, italic=False, code=False, link=None):
        if isinstance(node, NavigableString):
            text = str(node)
            if not text:
                return
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
            if code:
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xB4, 0x23, 0x18)
            if link:
                run.underline = True
                run.font.color.rgb = RGBColor(0x05, 0x4D, 0xC6)
            return

        if not isinstance(node, Tag):
            return

        tag = node.name.lower()

        if tag == 'br':
            paragraph.add_run('\n')
            return

        if tag in ('strong', 'b'):
            for child in node.children:
                add_inline(paragraph, child, bold=True, italic=italic, code=code, link=link)
            return

        if tag in ('em', 'i'):
            for child in node.children:
                add_inline(paragraph, child, bold=bold, italic=True, code=code, link=link)
            return

        if tag == 'code':
            for child in node.children:
                add_inline(paragraph, child, bold=bold, italic=italic, code=True, link=link)
            return

        if tag == 'a':
            href = node.get('href')
            for child in node.children:
                add_inline(paragraph, child, bold=bold, italic=italic, code=code, link=href)
            if href:
                paragraph.add_run(' (' + href + ')')
            return

        for child in node.children:
            add_inline(paragraph, child, bold=bold, italic=italic, code=code, link=link)

    def render_paragraph(tag):
        para = doc.add_paragraph()
        for child in tag.children:
            add_inline(para, child)

    def render_list(tag, ordered):
        style = 'List Number' if ordered else 'List Bullet'
        for li in tag.find_all('li', recursive=False):
            para = doc.add_paragraph(style=style)
            nested_lists = []
            for child in li.children:
                if isinstance(child, Tag) and child.name.lower() in ('ul', 'ol'):
                    nested_lists.append(child)
                else:
                    add_inline(para, child)

            for nested in nested_lists:
                render_list(nested, nested.name.lower() == 'ol')

    def render_table(tag):
        rows = tag.find_all('tr')
        if not rows:
            return

        max_cols = 0
        for row in rows:
            cols = 0
            for cell in row.find_all(['th', 'td']):
                try:
                    span = int(cell.get('colspan', 1))
                except (TypeError, ValueError):
                    span = 1
                cols += max(span, 1)
            max_cols = max(max_cols, cols)

        if max_cols <= 0:
            return

        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'

        for r_idx, row in enumerate(rows):
            c_idx = 0
            for cell in row.find_all(['th', 'td']):
                while c_idx < max_cols and table.rows[r_idx].cells[c_idx].text:
                    c_idx += 1
                if c_idx >= max_cols:
                    break

                target_cell = table.rows[r_idx].cells[c_idx]
                target_cell.text = ''
                p = target_cell.paragraphs[0]
                for child in cell.children:
                    add_inline(p, child)

                is_header = cell.name.lower() == 'th'
                if is_header:
                    for run in p.runs:
                        run.bold = True
                    set_cell_shading(target_cell, 'E9EEF5')

                try:
                    colspan = int(cell.get('colspan', 1))
                except (TypeError, ValueError):
                    colspan = 1

                colspan = max(colspan, 1)
                if colspan > 1 and c_idx + colspan - 1 < max_cols:
                    merged = target_cell.merge(table.rows[r_idx].cells[c_idx + colspan - 1])
                    merged.text = target_cell.text

                c_idx += colspan

        doc.add_paragraph('')

    def render_block(tag):
        name = tag.name.lower()

        if name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = min(max(int(name[1]), 1), 4)
            heading = doc.add_heading(level=level)
            for child in tag.children:
                add_inline(heading, child)
            return

        if name == 'p':
            render_paragraph(tag)
            return

        if name == 'pre':
            code_node = tag.find('code')
            code_text = code_node.get_text('\n', strip=False) if code_node else tag.get_text('\n', strip=False)
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            para.paragraph_format.left_indent = Pt(18)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(6)
            return

        if name == 'blockquote':
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(24)
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(5)
            for child in tag.children:
                add_inline(para, child, italic=True)
            return

        if name == 'ul':
            render_list(tag, ordered=False)
            return

        if name == 'ol':
            render_list(tag, ordered=True)
            return

        if name == 'table':
            render_table(tag)
            return

        if name == 'hr':
            para = doc.add_paragraph('-' * 40)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            return

    html_body = _markdown_to_html(content)
    soup = BeautifulSoup('<div>' + html_body + '</div>', 'html.parser')
    root = soup.find('div')

    if root:
        for node in root.children:
            if isinstance(node, Tag):
                render_block(node)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()
